import os
from sqlalchemy import create_engine
import docx
import re
import pymysql
from ExtractWordToMySQL.fields import match_list_v1, match_list_v2
from ExtractWordToMySQL.fields import fields_map_v1, fields_map_v2
from ExtractWordToMySQL.fields import host, database, user, password, port, provinces

# 初始化数据库连接
engine = create_engine('mysql+pymysql://root:123456@localhost:3306/process')
db = pymysql.connect(host, user, password, database, charset='utf8', port=port)
cursor = db.cursor()


#  存入数据库 字典格式  in (表名，字典）
def store_to_db(table_name, dic):
    keys = ','.join(dic.keys())
    values = ','.join(['%s'] * len(dic))
    sql = 'INSERT INTO {table}({keys}) VALUES ({values})'.format(table=table_name, keys=keys, values=values)
    cursor.execute(sql, tuple(dic.values()))
    db.commit()


#  提取表格内容，并存入数据库  in （docx中的表，报告编号，报告名称，flag=表类型）
def extra_from_table(table, doc_number, filename, flag):
    for row in table.rows[1:]:
        extra_result_dict = {}
        extra_result_dict['doc_name'] = filename
        extra_result_dict['doc_number'] = doc_number
        extra_result_dict['type'] = row.cells[0].text.replace('\n', '')
        extra_result_dict['number'] = row.cells[1].text
        extra_result_dict['assess_points'] = row.cells[2].text
        extra_result_dict['assessment'] = row.cells[3].text
        if len(row.cells) == 5:  # 2018多一个 评估记录 的情况
            extra_result_dict['assess_record'] = row.cells[4].text

        # for k, v in extra_result_dict.items():
        #     print(k + ':  ' + v)
        # print()

        # 将结果存入数据库
        if flag == 'risk':
            store_to_db('risk_analysis_table', extra_result_dict)
        elif flag == 'ensure':
            store_to_db('ensure_analysis_table', extra_result_dict)


#  处理 评估报告
def process_word_report(filepath, filename):
    # 读取文档
    try:
        docx_file = docx.Document(filepath)
    except:
        print('can`t open ' + filepath + ' ！！！')

    # 1、处理文档内容
    # 获取内容
    content = ''
    for para in docx_file.paragraphs:
        content = content + '\n' + para.text
    # 获取文档编号
    match = re.match('(.*)(报告编号：|报告编号:)(.*?)(\s)(.*)', content, re.S)
    if match is not None:
        doc_number = match.group(3).strip()
    # 存储 对应章节内容
    '''
    V1：2017，2018，其他 - 江西省一保通医联万家业务信息安全评估报告  网约车 - 江南出行平台业务信息安全评估报告  
    V2：网约车其余  
    V3：其他-“江西机关党建APP”安全评估报告
    '''
    flag = 2  # 修改匹配规则列表
    if flag == 1:
        match_list, fields_map = match_list_v1, fields_map_v1
    elif flag == 2:
        match_list, fields_map = match_list_v2, fields_map_v2

    for item in match_list:
        content_result_dict = {}
        content_result_dict['report_number'] = doc_number
        content_result_dict['report_name'] = filename
        content_result_dict['chapter'] = fields_map.get(item[0])[0]
        content_result_dict['chapter_number'] = fields_map.get(item[0])[1]

        # 根据规则调整
        if flag == 1:  # v1
            if item[0] == '(5.3\s?用户信息安全风险评估结果|5.3\s?业务平台安全风险评估结果)':
                if doc_number[0:4] == '2017':
                    content_result_dict['chapter'] = '用户信息安全风险评估结果'
                else:
                    content_result_dict['chapter'] = '业务平台安全风险评估结果'
            # chapter特殊情况2
            if item[0] == '(5.4\s?业务平台安全风险评估结果|5.4\s?业务运行安全风险评估结果)':
                if doc_number[0:4] == '2017':
                    content_result_dict['chapter'] = '业务平台安全风险评估结果'
                else:
                    content_result_dict['chapter'] = '业务运行安全风险评估结果'
            # 提取内容
            if item[0] == '(6\s?安全评估结论意见)':  # (2017、2018)  处理下结尾部分
                match = re.match('(.*)' + item[0] + '\s(.*?)(\s)', content, re.S)
            else:
                match = re.match('(.*)' + item[0] + '(.*)' + item[1] + '(.*)', content, re.S)

        elif flag == 2:  # v2
            # 提取内容
            if item[0] == '(7\s?整改情况及及复核结论)':  # 处理下结尾部分
                match = re.match('(.*)' + item[0] + '\s(.*?)(\s)', content, re.S)
            else:
                match = re.match('(.*)' + item[0] + '(.*)' + item[1] + '(.*)', content, re.S)

        if match is None:
            content_result_dict['content'] = ''
        else:
            content_result_dict['content'] = match.group(3).strip()

        # 查看存储结果 （调试）
        for k, v in content_result_dict.items():
            print(k + ': ' + v)
        print('------------------------------------')

        # 将结果存入数据库
        # store_to_db('assess_report_content', content_result_dict)

    # 2、存储 文档基本信息
    info_result_dict = {}
    info_result_dict['report_number'] = doc_number
    info_result_dict['report_name'] = filename
    info_result_dict['report_year'] = doc_number[0:4]
    for province in provinces:
        if province in filename:
            info_result_dict['report_province'] = province
    # 将结果存入数据库
    # store_to_db('assess_report_info', info_result_dict)

    # 查看存储结果 （调试）
    for k, v in info_result_dict.items():
        print(k + ': ' + v)
    print('------------------------------------')

    # 3、解析并存储 表格
    # 首先确定 风险分析表 risk_analysis_table 和 保障能力分析表 ensure_analysis_table
    for table in docx_file.tables:
        cell = table.rows[0].cells[0]
        if cell.text == '风险类型':
            risk_analysis_table = table
        elif cell.text == '保障能力类型':
            ensure_analysis_table = table
    # 提取并存储 风险分析表 和 保障能力分析表
    # extra_from_table(risk_analysis_table, doc_number, filename, flag='risk')
    # extra_from_table(ensure_analysis_table, doc_number, filename, flag='ensure')


if __name__ == '__main__':

    # 3、处理  江西评估报告
    rootpath = 'data/评估报告/待处理'
    for dirname in os.listdir(rootpath):
        dirpath = rootpath + '/' + dirname
        # print(dirpath)
        for filename in os.listdir(dirpath):
            filepath = dirpath + '/' + filename
            # print(filepath)
            # print(filename)
            # 默认是docx文件，如果是doc需要先转成docx
            process_word_report(filepath, filename[:-5])










