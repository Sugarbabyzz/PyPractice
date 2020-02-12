import os
import docx
import re
import pymysql
import datetime
from ExtractWordToMySQL.fields import host, database, user, password, port, chapter_number_map, provinces, years, operators

# 初始化数据库连接
db = pymysql.connect(host, user, password, database, charset='utf8', port=port)
cursor = db.cursor()


#  存入数据库 字典格式  in (表名，字典）
def store_to_db(table_name, dic):
    keys = ','.join(dic.keys())
    values = ','.join(['%s'] * len(dic))
    sql = 'INSERT INTO {table}({keys}) VALUES ({values})'.format(table=table_name, keys=keys, values=values)
    cursor.execute(sql, tuple(dic.values()))
    db.commit()


#  提取表格内容，并存入数据库  参数：（docx中的表对象，报告编号，报告名称，flag=表类型）
def extra_from_table(table, doc_number, filename, flag):
    for row in table.rows[1:]:
        try:
            extra_result_dict = {}
            extra_result_dict['report_name'] = filename
            extra_result_dict['report_number'] = doc_number
            extra_result_dict['type'] = row.cells[0].text.replace('\n', '')
            extra_result_dict['number'] = row.cells[1].text
            extra_result_dict['assess_points'] = row.cells[2].text
            extra_result_dict['assessment'] = row.cells[3].text
            if len(row.cells) == 5:  # 表格多一个 评估记录 的情况
                extra_result_dict['assess_record'] = row.cells[4].text

            # for k, v in extra_result_dict.items():
            #     print(k + ':  ' + v)
            # print()

            # 将结果存入数据库
            if flag == 'risk':
                store_to_db('risk_analysis_table', extra_result_dict)
            elif flag == 'ensure':
                store_to_db('ensure_analysis_table', extra_result_dict)
        except Exception as err:
            print('表格处理有错了：' + filename + ' ' + doc_number + ' ' + table + '\nerr:' + str(err))


#  解析 评估报告
def process_word_report(filepath, filename, province):
    # 读取文档
    try:
        # doc文件去读取对应docx文件
        if filename.endswith('.doc'):
            filepath = 'data/评估报告/对应docx/' + filename[:-4] + '.docx'
        docx_file = docx.Document(filepath)
    except Exception as err:
        print('can`t open ' + filepath + ' ！！！' + '\nerr' + str(err))

    # print(filepath + '----' + filename)
    # 处理文档内容
    # 1、解析 报告基本信息 （报告编号、省份、年份、业务名称、运营商(委托单位)）
    content = ''
    for para in docx_file.paragraphs[:20]:  # 读前20行即可获取基本信息
        content = content + '\n' + para.text
    # doc_number
    match = re.match('(.*)(报告编号：|报告编号:)\s*(.*?)(\s)(.*)', content, re.S)
    doc_number = match.group(3).strip() if match is not None else ''
    # year
    match = re.match('(.*)(\d{4})(\s*)年', content, re.S)
    year = match.group(2).strip() if match is not None else ''
    for y in years:
        if y in filename or y in doc_number:
            year = y
    # business
    match = re.match('(.*)“(.*)”(.*)', filename, re.S)
    business = match.group(2).strip() if match is not None else ''
    # operator
    operator = ''
    for opt in operators:
        if opt in filename:
            operator = opt
    match = re.match('(.*)(被评估单位：|被评估单位:|委托单位：|委托单位:|申请单位：|申请单位:)\s*(.*?)(\s)(.*)', content, re.S)
    if match is not None:
        operator = match.group(3).strip()
    # province
    # 从第一个表中解析基本信息
    table = docx_file.tables[0]
    for i in range(0, len(table.rows)):
        if '委托单位' in table.rows[i].cells[0].text or '被评估单位' in table.rows[i].cells[0].text or '被测评单位' in table.rows[i].cells[0].text or '受评单位' in table.rows[i].cells[0].text:
            if table.rows[i].cells[1].text != '':
                operator = table.rows[i].cells[1].text
        if '业务名称' in table.rows[i].cells[0].text:
            if table.rows[i].cells[1].text != '':
                business = table.rows[i].cells[1].text

    info_dict = {}
    info_dict['report_number'] = doc_number
    info_dict['report_name'] = filename
    info_dict['report_year'] = year
    info_dict['report_province'] = province
    info_dict['report_business'] = business
    info_dict['report_operator'] = operator

    # 查看存储结果 （调试）
    print('-——————————————————-——————————————————-——————————————————-——————————————————-——————————————————-———————————')
    print('-————-————' + filepath + '-————-————')  # 打印当前处理报告的路径
    print('-————-————' + datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S') + '-————-————')  # 打印当前时间

    # 2、解析并存储 报告子章节内容
    #   content_dict 存储 <标题-标题下内容>
    flag = 0
    content = ''
    content_dict = {}
    chapter = '0'
    for p in docx_file.paragraphs:
        if p.style.name is None:
            continue
        if p.style.name == 'Heading 1' or p.style.name == 'Heading 2' or '标题' in p.style.name:
            if flag == 0:
                flag = 1
            else:
                content_dict[chapter] = content
                content = ''
            chapter = p.text
        else:
            if p.style.name != 'Heading 1' and p.style.name != 'Heading 2' and '标题' not in p.style.name and flag == 1:
                content = content + '\n' + p.text
    # 处理一下全篇没有标题的情况
    if chapter == '0':
        for p in docx_file.paragraphs[16:]:
            content = content + '\n' + p.text
        content_dict[chapter] = content
    else:
        content_dict[chapter] = content

    # 处理章节和内容字段，存入数据库
    for k, v in content_dict.items():
        # print('--------------------------------------------')
        result_dict = {}
        result_dict['report_number'] = doc_number
        result_dict['report_name'] = filename
        result_dict['chapter'] = ''.join(re.findall('[^\d?\.?\d?\.?\d?\d*]', k)).strip() if re.findall('[^\d?\.?\d?\.?\d?\d*]', k) else ''
        result_dict['chapter_number'] = re.findall('\d?\.?\d?\.?\d?\d*', k)[0].strip() if re.findall('\d?\.?\d?\.?\d?\d*', k) else ''
        for key, value in chapter_number_map.items():
            if key in k:
                result_dict['chapter_number'] = chapter_number_map.get(key)
        result_dict['content'] = v.strip()

        # print('----- ' + result_dict['chapter_number'] + ' - ' + result_dict['chapter'] + ' -----')
        # for key, value in result_dict.items():
        #     print(key + ': ' + value)

        # 存 业务名称 基本信息
        if result_dict['chapter'] == '业务名称':
            info_dict['report_business'] = re.sub('[^\w\u4e00-\u9fff()（）]+', '', result_dict['content'])

        # 将结果存入数据库
        # store_to_db('assess_report_content', result_dict)

    # 3、解析并存储 表格
    # 首先确定 风险分析表 risk_analysis_table 和 保障能力分析表 ensure_analysis_table
    # risk_analysis_table, ensure_analysis_table = '', ''
    # for table in docx_file.tables:
    #     cell = table.rows[0].cells[0]
    #     if cell.text == '风险类型':
    #         risk_analysis_table = table
    #     elif cell.text == '保障能力类型':
    #         ensure_analysis_table = table
    # # 提取并存储 风险分析表 和 保障能力分析表
    # if risk_analysis_table != '':
    #     extra_from_table(risk_analysis_table, doc_number, filename, flag='risk')
    # if ensure_analysis_table != '':
    #     extra_from_table(ensure_analysis_table, doc_number, filename, flag='ensure')

    # 4、存储 报告基本信息
    store_to_db('assess_report_info', info_dict)

    for k, v in info_dict.items():
        print(k + ': ' + v)


# 主程序
if __name__ == '__main__':

    # 处理  评估报告
    rootpath = 'data/评估报告/处理中'
    error_files = []
    processed_files = []
    for dirname in os.listdir(rootpath):
        if dirname == '.DS_Store':
            continue
        dirpath = rootpath + '/' + dirname
        # dirname中一定要包含报告的省份
        for filename in os.listdir(dirpath):
            if filename == '.DS_Store':
                continue
            filepath = dirpath + '/' + filename
            # 默认是docx文件，如果是doc需要先转成docx（这里从<对应docx>文件夹下找）
            # 调用word解析程序，参数：（路径，文件名，省份）
            try:
                process_word_report(filepath, filename, dirname.replace('评估报告', ''))
            except Exception as err:
                print('处理错误！ ： ' + filepath)
                error_files.append(filepath + '\n错误日志：' + str(err))
            # 加入已处理过的文件日志LOG
            processed_files.append(datetime.datetime.now().strftime(
                '%Y-%m-%d %H:%M:%S') + '：' + dirname + ' - ' + filename + ' Done！')

        # 打印日志到txt
        with open('processing_log.txt', 'a') as f:
            for file in processed_files:
                f.write(file + '\n')
            processed_files = []

    # 打印处理出错的文件日志到txt
    with open('error_files.txt', 'a') as file_object:
        for file in error_files:
            file_object.write(file + '\n\n')










