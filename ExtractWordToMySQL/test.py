import docx
import pandas as pd
import re
import os


# check 标题格式是否符合要求
# cncert
rootpath = 'data/评估报告处理中/待处理'
for dirname in os.listdir(rootpath):
    dirpath = rootpath + '/' + dirname
    for filename in os.listdir(dirpath):
        print('-——————————————————-——————————————————-——————————————————')

        filepath = dirpath + '/' + filename
        print('-————-————' + filepath + '-————-————')
        # 默认是docx文件，如果是doc需要先转成docx
        docx_file = docx.Document(filepath)
        for p in docx_file.paragraphs:
            if p.style.name == 'Heading 1' or p.style.name == 'Heading 2' or '标题' in p.style.name:
                print(p.style.name + ' : ' + p.text.strip())


# flag = 0
# content = ''
# dic = {}
# for p in docx_file.paragraphs:
#     # if re.match('^Heading \d+$', p.style.name):
#     if p.style.name == 'Heading 2' or p.style.name == 'Heading 1':
#         if flag == 0:
#             flag = 1
#         else:
#         # print(chapter + ': ' + content)
#             dic[chapter] = content
#             content = ''
#         chapter = p.text
#     else:
#         if p.style.name == 'Normal' and flag == 1:
#             content = content + '\n' + p.text
# dic[chapter] = content
#
# for k, v in dic.items():
#     print(k + v)
#     print('-----––––––-----––––––-----––––––-----––––––')

