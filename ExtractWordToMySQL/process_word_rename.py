import os
import docx
import re
import datetime
from ExtractWordToMySQL.fields import chapter_number_map, years, operators


#  解析 评估报告
def process_word_report(filepath, filename, dirname, dirpath):
    # 读取文档
    old_filepath = filepath
    try:
        # doc文件去读取对应docx文件
        if filename.endswith('.doc'):
            filepath = 'data/评估报告/第一批对应docx/' + filename[:-4] + '.docx'
        docx_file = docx.Document(filepath)
    except Exception as err:
        print('can`t open ' + filepath + ' ！！！' + '\nerr' + str(err))

    # 1、解析 报告基本信息 （报告编号、省份、年份、业务名称、厂商单位）
    content = ''
    for para in docx_file.paragraphs[:20]:  # 读前20行即可获取基本信息
        content = content + '\n' + para.text
    # doc_number
    match = re.match('(.*)(报告编号：|报告编号:)\s*(.*?)(\s)(.*)', content, re.S)
    doc_number = match.group(3).strip() if match is not None else ''
    # year
    match = re.match('(.*)(\d{4})(\s*)年', content, re.S)
    year = match.group(2).strip() if match is not None else ''
    for y in years:
        if y in filename or y in doc_number:
            year = y
    # business
    match = re.match('(.*)“(.*)”(.*)', filename, re.S)
    business = match.group(2).strip() if match is not None else ''
    match = re.match('(.*)(业务名称：|业务名称:|系统名称：|系统名称:|产品名称：|产品名称:)\s*(.*?)(\s)(.*)', content, re.S)
    if match is not None:
        business = match.group(3).strip()
    # operator
    operator = ''
    for opt in operators:
        if opt in filename:
            operator = opt
    match = re.match('(.*)(被评估单位：|被评估单位:|委托单位：|委托单位:|申请单位：|申请单位:)\s*(.*?)(\s)(.*)', content, re.S)
    if match is not None:
        operator = match.group(3).strip()
    match = re.match('(.*)中国(.*?)公司(.*)', content, re.S)
    if match is not None:
        operator = '中国' + match.group(2).strip() + '公司'

    # 从第一个表中解析基本信息
    if docx_file.tables:
        table = docx_file.tables[0]
        for i in range(0, len(table.rows)):
            if '委托单位' in table.rows[i].cells[0].text or '被评估单位' in table.rows[i].cells[0].text or '被测评单位' in table.rows[i].cells[0].text or '受评单位' in table.rows[i].cells[0].text:
                if table.rows[i].cells[1].text.strip() != '':
                    operator = table.rows[i].cells[1].text.split('\n')[0].strip()
            if '业务名称' in table.rows[i].cells[0].text:
                if table.rows[i].cells[1].text.strip() != '':
                    business = table.rows[i].cells[1].text.split('\n')[0].strip()

    info_dict = {}
    info_dict['report_number'] = doc_number
    info_dict['report_name'] = filename
    info_dict['report_year'] = year
    info_dict['report_province'] = dirname.replace('评估报告', '')
    info_dict['report_business'] = business
    info_dict['report_operator'] = operator

    # 查看存储结果 （调试）
    print('-——————————————————-——————————————————-——————————————————-——————————————————-——————————————————-———————————')
    print('-————-————' + filepath + '-————-————')  # 打印当前处理报告的路径
    print('-————-————' + datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S') + '-————-————')  # 打印当前时间

    # 2、解析并存储 报告子章节内容
    #   content_dict 存储 <标题-标题下内容>
    flag = 0
    content = ''
    content_dict = {}
    chapter = '0'
    for p in docx_file.paragraphs:
        if p.style.name is None:
            continue
        if p.style.name == 'Heading 1' or p.style.name == 'Heading 2' or '标题' in p.style.name:
            if flag == 0:
                flag = 1
            else:
                content_dict[chapter] = content
                content = ''
            chapter = p.text
        else:
            if p.style.name != 'Heading 1' and p.style.name != 'Heading 2' and '标题' not in p.style.name and flag == 1:
                content = content + '\n' + p.text
    # 处理一下全篇没有标题的情况
    if chapter == '0':
        for p in docx_file.paragraphs[16:]:
            content = content + '\n' + p.text
        content_dict[chapter] = content
    else:
        content_dict[chapter] = content

    # 处理章节和内容字段，存入数据库
    for k, v in content_dict.items():
        result_dict = {}
        result_dict['report_number'] = doc_number
        result_dict['report_name'] = filename
        result_dict['chapter'] = ''.join(re.findall('[^\d?\.?\d?\.?\d?\d*]', k)).strip() if re.findall('[^\d?\.?\d?\.?\d?\d*]', k) else ''
        result_dict['chapter_number'] = re.findall('\d?\.?\d?\.?\d?\d*', k)[0].strip() if re.findall('\d?\.?\d?\.?\d?\d*', k) else ''
        for key, value in chapter_number_map.items():
            if key in k:
                result_dict['chapter_number'] = chapter_number_map.get(key)
        result_dict['content'] = v.strip()

        # 存 业务名称 基本信息
        if result_dict['chapter'] == '业务基本情况介绍' and business == '':
            match = re.match('(.*?)“(.*?)”(.*)', result_dict['content'], re.S)
            if match is not None:
                business = match.group(2).strip()

        if result_dict['chapter'] == '业务名称' and business == '':
            business = re.sub('[^\w\u4e00-\u9fff()（）]+', '', result_dict['content'])

        if result_dict['chapter'] == '评估对象' and business == '':
            business = re.sub('[^\w\u4e00-\u9fff()（）]+', '', result_dict['content'])

    # 融媒体中心的报告，从第一张表格中解析业务名称
    if docx_file.tables and business == '':
        table = docx_file.tables[0]
        if table.columns[0].cells[0].text.strip() == '被检查企业':
            if table.columns[3].cells[0].text.strip() != '':
                business = table.columns[3].cells[0].text.strip()

        if business == '':
            for i in range(0, len(table.columns)):
                c = table.columns[i].cells[0].text.strip()
                match = re.match('(.*?)“(.*?)”(.*)', c, re.S)
                if match is not None:
                    business = match.group(2).strip()
                    break

        # 如果业务名称没有被引号括起来
        if business == '' and len(docx_file.tables) > 1:
            table = docx_file.tables[1]
            if table.columns[2].cells[0].text.strip() != '':
                business = table.columns[2].cells[0].text.strip()

    # 确定最终的业务名称
    info_dict['report_business'] = business

    for k, v in info_dict.items():
        print(k + ': ' + v)

    # 重命名
    new_filename = '[' + info_dict['report_year'] + ']' + '[' + info_dict['report_operator'] +\
                   ']' + '[' + info_dict['report_business'] + ']新技术新业务信息安全评估报告'
    new_filepath = dirpath + '/' + new_filename + '.' + filename.split('.')[-1]
    print('-————-————' + '重命名后的文件名：' + new_filename + '.' + filename.split('.')[-1] + '-————-————')
    os.rename(old_filepath, new_filepath)

    # 打印有问题的文件日志到txt
    if info_dict['report_year'] == '2009' or info_dict['report_year'] == '' or info_dict['report_operator'] == '' or info_dict['report_business'] == '':
        with open('something_wrong_files.txt', 'a') as file_object:
            file_object.write(new_filepath + '\n')


# 主程序
if __name__ == '__main__':

    # 处理  评估报告
    rootpath = 'data/评估报告/处理中'
    error_files = []
    processed_files = []
    for dirname in os.listdir(rootpath):
        if dirname == '.DS_Store':
            continue
        dirpath = rootpath + '/' + dirname
        # dirname中一定要包含报告的省份
        for filename in os.listdir(dirpath):
            if filename == '.DS_Store':
                continue
            filepath = dirpath + '/' + filename
            # 默认是docx文件，如果是doc需要先转成docx（这里从<对应docx>文件夹下找）
            # 调用word解析程序，参数：（路径，文件名，省份）
            try:
                process_word_report(filepath, filename, dirname, dirpath)
            except Exception as err:
                print('处理错误！ ： ' + filepath)
                error_files.append(filepath + '\n错误日志：' + str(err))
            # 加入已处理过的文件日志LOG
            processed_files.append(datetime.datetime.now().strftime(
                '%Y-%m-%d %H:%M:%S') + '：' + dirname + ' - ' + filename + ' Done！')

        # 打印日志到txt
        with open('processing_log.txt', 'a') as f:
            for file in processed_files:
                f.write(file + '\n')
            processed_files = []

    # 打印处理出错的文件日志到txt
    with open('error_files.txt', 'a') as file_object:
        for file in error_files:
            file_object.write(file + '\n\n')










