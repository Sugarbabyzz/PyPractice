import pandas as pd
import os
from sqlalchemy import create_engine
import docx
import re
import pymysql
import subprocess
from ProcessDataToMySQL.fields import match_list
from ProcessDataToMySQL.fields import fields_map
from ProcessDataToMySQL.fields import host, database, user, password, port, years, provinces


# 初始化数据库连接
engine = create_engine('mysql+pymysql://root:123456@localhost:3306/process')
db = pymysql.connect(host, user, password, database, charset='utf8', port=port)
cursor = db.cursor()


#  处理 评估专家推荐表（新疆）.xlsx
def process_excel_xj(filepath):
    print('评估专家推荐表（新疆）.xlsx  Processing')
    columns = ['name', 'province', 'recom_level', 'more_than_twice', 'cisp_number', 'other_cert', 'work_years', 'recom_reason']
    data = pd.read_excel(filepath, names=columns)
    data.to_sql('pgzjtjb_xj', con=engine, if_exists='append', index=False)
    print('评估专家推荐表（新疆）.xlsx  Done')


#  处理 评估报告登记表（重庆）.xlsx
def process_excel_cq(filepath):
    columns = ['assess_report_time', 'assess_report_title', 'category', 'intro', 'user_scale', 'risk_points', 'assessor',
               'project_province', 'assessor_province', 'counterpart']
    data = pd.read_excel(filepath, names=columns)
    data['assess_report_time'] = data.loc[:, 'assess_report_time'].apply(lambda x: x.strftime('%Y-%m-%d'))
    data.to_sql('pgbgdjb_cq', con=engine, if_exists='append', index=False)


#  处理 评估报告
def process_word_report(filepath, filename, dirname):
    # 读取文档
    try:
        docx_file = docx.Document(filepath)
    except:
        print('can`t open ' + filepath + ' ！！！')

    # 1、处理文档内容
    #   获取内容
    docx_para = docx_file.paragraphs
    content = ''  # 存储word中的文字内容
    for para in docx_para:
        content = content + '\n' + para.text
    # print(content)
    #   提取文档中子标题下的内容
    result_dict = {}
    result_dict['report_name'] = filename
    for item in match_list:
        # 暂不处理 附件
        if item[0] == '(报告编号：|报告编号:)':
            match = re.match('(.*)' + item[0] + '(.*?)' + item[1] + '(.*)', content, re.S)
        elif item[0] == '(6\s?安全评估结论意见)':
            match = re.match('(.*)' + item[0] + '\s(.*?)(\s)', content, re.S)
        else:
            match = re.match('(.*)' + item[0] + '(.*)' + item[1] + '(.*)', content, re.S)
        # 解析的内容存入到对应字段中
        if match is None:
            result_dict[fields_map.get(item[0])] = ''
        else:
            result_dict[fields_map.get(item[0])] = match.group(3).strip()
    result_dict['report_year'] = result_dict.get('report_number')[0:4]
    for province in provinces:
        if province in filename:
            result_dict['report_province'] = province

    # 查看存储结果 （调试）
    for k, v in result_dict.items():
        print(k + ': \n' + v)
        print('------------------------------------')

    # 将结果存入数据库
    table = 'assess_report'
    keys = ','.join(result_dict.keys())
    values = ','.join(['%s'] * len(result_dict))
    sql = 'INSERT INTO {table}({keys}) VALUES ({values})'.format(table=table, keys=keys, values=values)
    cursor.execute(sql, tuple(result_dict.values()))
    db.commit()

    # 2、处理表格
    # 首先确定 风险分析表 risk_analysis_table 和 保障能力分析表 ensure_analysis_table
    for table in docx_file.tables:
        row = table.rows[0]
        cell = row.cells[0]
        if cell.text == '风险类型':
            risk_analysis_table = table
        elif cell.text == '保障能力类型':
            ensure_analysis_table = table
    print('--------')
    # 提取 风险分析表
    risk_result_dict = {}
    for row in risk_analysis_table.rows[1:]:
        risk_result_dict['doc_name'] = filename
        risk_result_dict['doc_number'] = result_dict.get('report_number')
        risk_result_dict['risk_type'] = row.cells[0].text.replace('\n', '')
        risk_result_dict['risk_number'] = row.cells[1].text
        risk_result_dict['risk_assess_points'] = row.cells[2].text
        risk_result_dict['assessment'] = row.cells[3].text
        if len(row.cells) == 5:  # 2018多一个 评估记录 的情况
            risk_result_dict['assess_record'] = row.cells[4].text

        for k, v in risk_result_dict.items():
            print(k + ':  ' + v)
        # 将结果存入数据库
        risk_table = 'risk_analysis_table'
        keys = ','.join(risk_result_dict.keys())
        values = ','.join(['%s'] * len(risk_result_dict))
        sql = 'INSERT INTO {table}({keys}) VALUES ({values})'.format(table=risk_table, keys=keys, values=values)
        cursor.execute(sql, tuple(risk_result_dict.values()))
        db.commit()

    # 提取 保障能力分析表
    ensure_result_dict = {}
    for row in ensure_analysis_table.rows[1:]:
        ensure_result_dict['doc_name'] = filename
        ensure_result_dict['doc_number'] = result_dict.get('report_number')
        ensure_result_dict['ensure_type'] = row.cells[0].text.replace('\n', '')
        ensure_result_dict['ensure_number'] = row.cells[1].text
        ensure_result_dict['ensure_assess_points'] = row.cells[2].text
        ensure_result_dict['assessment'] = row.cells[3].text
        if len(row.cells) == 5:  # 2018多一个 评估记录 的情况
            ensure_result_dict['assess_record'] = row.cells[4].text

        for k, v in ensure_result_dict.items():
            print(k + ':  ' + v)
        print()
        # 将结果存入数据库
        ensure_table = 'ensure_analysis_table'
        keys = ','.join(ensure_result_dict.keys())
        values = ','.join(['%s'] * len(ensure_result_dict))
        sql = 'INSERT INTO {table}({keys}) VALUES ({values})'.format(table=ensure_table, keys=keys, values=values)
        cursor.execute(sql, tuple(ensure_result_dict.values()))
        db.commit()


if __name__ == '__main__':

    # # 1、处理  评估专家推荐表（新疆）.xlsx
    # filepath_xj = 'data/评估专家推荐表（新疆）.xlsx'
    # process_excel_xj(filepath_xj)
    #
    # # 2、处理  评估报告登记表（重庆）.xlsx
    # filepath_cq = 'data/评估报告登记表（重庆）.xlsx'
    # process_excel_cq(filepath_cq)

    # 3、处理  江西评估报告
    rootpath = 'data/评估报告'
    for dirname in os.listdir(rootpath):
        dirpath = rootpath + '/' + dirname
        print(dirpath)
        for filename in os.listdir(dirpath):
            filepath = dirpath + '/' + filename
            print(filepath)
            print(filename)
            process_word_report(filepath, filename[:-5], dirname)








