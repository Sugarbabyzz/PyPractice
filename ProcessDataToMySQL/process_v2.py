import pandas as pd
import os
from sqlalchemy import create_engine
import docx
import re
import pymysql
from ProcessDataToMySQL.fields import match_list_v1, match_list_v2
from ProcessDataToMySQL.fields import fields_map_v1, fields_map_v2
from ProcessDataToMySQL.fields import host, database, user, password, port, provinces

# 初始化数据库连接
engine = create_engine('mysql+pymysql://root:123456@localhost:3306/process')
db = pymysql.connect(host, user, password, database, charset='utf8', port=port)
cursor = db.cursor()


#  处理 评估专家推荐表（新疆）.xlsx
def process_expert_recom_table(filepath):
    columns = ['name', 'province', 'recom_level', 'more_than_twice', 'cisp_number', 'other_cert', 'work_years',
               'recom_reason']
    data = pd.read_excel(filepath, names=columns)
    data.to_sql('expert_recom_table', con=engine, if_exists='append', index=False)


#  处理 评估报告登记表（重庆）.xlsx
def process_report_register_table(filepath):
    columns = ['assess_time', 'report_name', 'category', 'intro', 'user_scale', 'risk_points',
               'assessor', 'project_province', 'assessor_province', 'counterpart']
    data = pd.read_excel(filepath, names=columns)
    data['assess_time'] = data.loc[:, 'assess_time'].apply(lambda x: x.strftime('%Y-%m-%d'))
    data.to_sql('report_register_table', con=engine, if_exists='append', index=False)


#  存入数据库 字典格式  in (表名，字典）
def store_to_db(table_name, dic):
    keys = ','.join(dic.keys())
    values = ','.join(['%s'] * len(dic))
    sql = 'INSERT INTO {table}({keys}) VALUES ({values})'.format(table=table_name, keys=keys, values=values)
    cursor.execute(sql, tuple(dic.values()))
    db.commit()


#  提取表格内容，并存入数据库  in （docx中的表，报告编号，报告名称，flag=表类型）
def extra_from_table(table, doc_number, filename, flag):
    for row in table.rows[1:]:
        extra_result_dict = {}
        extra_result_dict['doc_name'] = filename
        extra_result_dict['doc_number'] = doc_number
        extra_result_dict['type'] = row.cells[0].text.replace('\n', '')
        extra_result_dict['number'] = row.cells[1].text
        extra_result_dict['assess_points'] = row.cells[2].text
        extra_result_dict['assessment'] = row.cells[3].text
        if len(row.cells) == 5:  # 2018多一个 评估记录 的情况
            extra_result_dict['assess_record'] = row.cells[4].text

        # for k, v in extra_result_dict.items():
        #     print(k + ':  ' + v)
        # print()

        # 将结果存入数据库
        if flag == 'risk':
            store_to_db('risk_analysis_table', extra_result_dict)
        elif flag == 'ensure':
            store_to_db('ensure_analysis_table', extra_result_dict)


#  处理 评估报告
def process_word_report(filepath, filename):
    # 读取文档
    try:
        docx_file = docx.Document(filepath)
    except:
        print('can`t open ' + filepath + ' ！！！')

    # 1、处理文档内容
    # 获取内容
    content = ''
    for para in docx_file.paragraphs:
        content = content + '\n' + para.text
    # 获取文档编号
    match = re.match('(.*)(报告编号：|报告编号:)(.*?)(\s)(.*)', content, re.S)
    if match is not None:
        doc_number = match.group(3).strip()
    # 存储 对应章节内容
    for item in match_list_v1:
        content_result_dict = {}
        content_result_dict['report_number'] = doc_number
        content_result_dict['report_name'] = filename
        content_result_dict['chapter'] = fields_map_v1.get(item[0])[0]
        content_result_dict['chapter_number'] = fields_map_v1.get(item[0])[1]

        # chapter特殊情况1
        if item[0] == '(5.3\s?用户信息安全风险评估结果|5.3\s?业务平台安全风险评估结果)':
            if doc_number[0:4] == '2017':
                content_result_dict['chapter'] = '用户信息安全风险评估结果'
            else:
                content_result_dict['chapter'] = '业务平台安全风险评估结果'
        # chapter特殊情况2
        if item[0] == '(5.4\s?业务平台安全风险评估结果|5.4\s?业务运行安全风险评估结果)':
            if doc_number[0:4] == '2017':
                content_result_dict['chapter'] = '业务平台安全风险评估结果'
            else:
                content_result_dict['chapter'] = '业务运行安全风险评估结果'
        # 提取内容
        if item[0] == '(6\s?安全评估结论意见)':  # (2017、2018)  处理下结尾部分
            match = re.match('(.*)' + item[0] + '\s(.*?)(\s)', content, re.S)
        else:
            match = re.match('(.*)' + item[0] + '(.*)' + item[1] + '(.*)', content, re.S)

        if match is None:
            content_result_dict['content'] = ''
        else:
            content_result_dict['content'] = match.group(3).strip()

        # 查看存储结果 （调试）
        for k, v in content_result_dict.items():
            print(k + ': ' + v)
        print('------------------------------------')

        # 将结果存入数据库
        store_to_db('assess_report_content', content_result_dict)

    # 存储 文档基本信息
    info_result_dict = {}
    info_result_dict['report_number'] = doc_number
    info_result_dict['report_name'] = filename
    info_result_dict['report_year'] = doc_number[0:4]
    for province in provinces:
        if province in filename:
            info_result_dict['report_province'] = province
    store_to_db('assess_report_info', info_result_dict)

    # 查看存储结果 （调试）
    for k, v in info_result_dict.items():
        print(k + ': ' + v)
    print('------------------------------------')

    # 2、处理表格
    # 首先确定 风险分析表 risk_analysis_table 和 保障能力分析表 ensure_analysis_table
    for table in docx_file.tables:
        cell = table.rows[0].cells[0]
        if cell.text == '风险类型':
            risk_analysis_table = table
        elif cell.text == '保障能力类型':
            ensure_analysis_table = table
    # 提取 风险分析表
    extra_from_table(risk_analysis_table, doc_number, filename, flag='risk')
    # 提取 保障能力分析表
    extra_from_table(ensure_analysis_table, doc_number, filename, flag='ensure')


if __name__ == '__main__':

    # 1、处理  评估专家推荐表（新疆）.xlsx
    filepath_xj = 'data/评估专家推荐表（新疆）.xlsx'
    process_expert_recom_table(filepath_xj)

    # 2、处理  评估报告登记表（重庆）.xlsx
    filepath_cq = 'data/评估报告登记表（重庆）.xlsx'
    process_report_register_table(filepath_cq)

    # 3、处理  江西评估报告
    # rootpath = 'data/评估报告'
    # for dirname in os.listdir(rootpath):
    #     dirpath = rootpath + '/' + dirname
    #     # print(dirpath)
    #     for filename in os.listdir(dirpath):
    #         filepath = dirpath + '/' + filename
    #         # print(filepath)
    #         # print(filename)
    #         # 默认是docx文件，如果是doc需要先转成docx
    #         process_word_report(filepath, filename[:-5])










